r"""Résumé parser — file in (pdf/docx/txt/image), structured JSON out.

Pipeline: extract the résumé to plain text, hand that text to `claude -p`, and
return the strict JSON the model produces (the talent-radar schema).

It is read-only and side-effect-free (ADR-0001): it reads one file and returns data.

Dependencies (install only what your input formats need):
    uv add pypdf python-docx        # pdf / docx
Images need nothing extra — they are read by Claude's own vision (no Tesseract).
`claude` (Claude Code CLI) must be on PATH for every format.

Usage:
    python -m tools.resume.parser <resume_file> [--role "Backend Engineer"] [--out out.json]
"""

from __future__ import annotations

import argparse
import json
import shutil
import subprocess
import sys
from pathlib import Path

from brain.engine import run_llm

# The extraction contract handed to Claude. {{TARGET_ROLE}} / {{RESUME_TEXT}} are
# filled by str.replace (NOT .format — the prompt is full of literal JSON braces).
PROMPT = r'''You are a résumé-parsing engine for a hiring/talent system. You receive the raw text of ONE
candidate's résumé and return a single, strictly-valid JSON object that distils the résumé into
structured fields plus a 6-axis capability score. Your output is consumed by software, not a human —
output JSON only.

==================== HARD RULES ====================
1. Output ONE JSON object and NOTHING else — no markdown, no code fences, no text before/after.
2. Use ONLY information supported by the résumé text. Never invent facts, employers, dates, or skills.
3. If a field is absent, set it to null (or [] for lists) AND add its name to "missing_fields". Do not guess.
4. Every radar "evidence" string MUST cite a specific detail from the résumé that justifies the score.
   If there is no evidence for an axis, score it low and say so in evidence.
5. Score conservatively and use the RUBRIC below so scores are comparable across candidates.
   Scores are integers 0–100.
6. "salary": fill ONLY if the résumé explicitly states a current or expected salary. Otherwise null.
   Never estimate compensation.
7. Lower "parse_confidence" (0–1) when the résumé is sparse/messy or you had to infer a lot. Put
   anything suspicious (gaps, inconsistencies, unverifiable claims) in "flags".

==================== RADAR — score each axis 0–100 using these anchors ====================
Technical:
- technical_depth — proficiency/seniority in their core craft.
    20 = little technical evidence · 50 = solid mid-level in their main stack ·
    80 = deep expert doing advanced/complex work · 100 = rare top-tier authority.
- stack_breadth — range of tools/languages/technologies.
    20 = one narrow tool · 50 = a few related technologies · 80 = wide, polyglot across areas.
- delivery_scale — complexity & scale of what they have ACTUALLY shipped.
    20 = only coursework/personal toys · 50 = shipped real features/projects ·
    80 = led large-scale, high-impact, production systems.
Non-technical:
- work_experience — years + role progression + seniority.
    20 = <1 yr / intern · 50 = ~3–5 yrs steady · 80 = 8+ yrs with clear progression to senior/lead.
- communication — writing, docs, presentations, client/stakeholder & cross-functional work
  (inferred from roles + résumé quality).
    20 = minimal signal · 50 = some collaborative/cross-functional roles ·
    80 = strong evidence (client-facing, led presentations, wrote docs, taught/mentored).
- leadership_ownership — leading people and owning initiatives end-to-end.
    20 = individual contributor, no ownership signal · 50 = owned projects/small initiatives ·
    80 = led teams and drove org-level outcomes.

==================== OUTPUT SCHEMA (fill every key; null/[] when unknown) ====================
{
  "schema_version": "1.0",
  "parse_confidence": 0.0,
  "name": null,
  "email": null,
  "phone": null,
  "location": null,
  "links": { "linkedin": null, "github": null, "portfolio": null },
  "headline": null,
  "summary": null,
  "total_years_experience": null,
  "seniority": null,
  "skills": [],
  "strengths": [],
  "domains": [],
  "roles": [
    { "company": null, "title": null, "start": null, "end": null, "months": null, "highlights": [] }
  ],
  "education": [
    { "degree": null, "field": null, "institution": null, "year": null }
  ],
  "certifications": [],
  "achievements": [],
  "radar": {
    "technical_depth":      { "score": 0, "evidence": null },
    "stack_breadth":        { "score": 0, "evidence": null },
    "delivery_scale":       { "score": 0, "evidence": null },
    "work_experience":      { "score": 0, "evidence": null },
    "communication":        { "score": 0, "evidence": null },
    "leadership_ownership": { "score": 0, "evidence": null }
  },
  "salary": null,
  "missing_fields": [],
  "flags": []
}

Field notes:
- "summary": 2–3 sentence "who they are / what they're good at."
- "strengths": short list of what this person is clearly good at.
- "seniority": one of "junior" | "mid" | "senior" | "lead" | "principal".
- "skills": concrete skills/technologies as short strings.
- "roles": most recent first; "months" = duration in months if derivable.
- "salary" (only if stated): { "amount": number, "currency": "ISO e.g. USD",
    "basis": "annual" | "monthly", "kind": "current" | "expected" }.

==================== INPUT ====================
TARGET ROLE (optional, may be blank): {{TARGET_ROLE}}
If a role is given, weight technical_depth / stack_breadth / delivery_scale toward skills relevant to
that role; if blank, judge in general terms.

RÉSUMÉ TEXT:
"""
{{RESUME_TEXT}}
"""

Return the JSON object now.'''

IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tif", ".tiff"}


def _missing(pkg: str, suffix: str) -> SystemExit:
    return SystemExit(f"Reading {suffix} files needs '{pkg}'. Install it: uv add {pkg}")


def extract_text(path: Path) -> str:
    """Dispatch on file extension and return the résumé as plain text."""
    suffix = path.suffix.lower()

    if suffix in {".txt", ".md"}:
        return path.read_text(encoding="utf-8", errors="replace")

    if suffix == ".pdf":
        try:
            from pypdf import PdfReader
        except ImportError:
            raise _missing("pypdf", suffix)
        reader = PdfReader(str(path))
        return "\n".join(page.extract_text() or "" for page in reader.pages)

    if suffix == ".docx":
        try:
            import docx
        except ImportError:
            raise _missing("python-docx", suffix)
        doc = docx.Document(str(path))
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                parts.extend(cell.text for cell in row.cells)
        return "\n".join(parts)

    if suffix in IMAGE_SUFFIXES:
        return _ocr_image(path)

    raise SystemExit(
        f"Unsupported file type '{suffix}'. Supported: .pdf .docx .txt .md "
        f"and images ({', '.join(sorted(IMAGE_SUFFIXES))})."
    )


def _ocr_image(path: Path) -> str:
    """Transcribe a résumé image to text using Claude's vision (no Tesseract needed)."""
    claude = shutil.which("claude")
    if not claude:
        raise SystemExit("`claude` CLI not found on PATH. Install Claude Code and retry.")
    ref = path.resolve().as_posix()
    prompt = (
        "Transcribe ALL text from this résumé image exactly as written, preserving line "
        f"order. Output only the transcribed text, no commentary. @{ref}"
    )
    proc = subprocess.run(
        [claude, "-p", prompt],
        capture_output=True,
        text=True,
        encoding="utf-8",
        timeout=300,
    )
    if proc.returncode != 0:
        raise SystemExit(f"claude image read failed (exit {proc.returncode}):\n{proc.stderr.strip()}")
    return proc.stdout.strip()


def _strip_fences(text: str) -> str:
    """Defensively remove ```json ... ``` fences if the model adds them."""
    if text.startswith("```"):
        lines = text.splitlines()
        lines = lines[1:] if lines and lines[0].startswith("```") else lines
        if lines and lines[-1].strip() == "```":
            lines = lines[:-1]
        text = "\n".join(lines)
    return text.strip()


def parse_resume(path: Path, role: str = "") -> dict:
    """Extract text, run Claude, and return the parsed JSON object."""
    text = extract_text(path).strip()
    if not text:
        raise SystemExit(f"No text could be extracted from {path} (empty or image needs OCR).")
    prompt = PROMPT.replace("{{TARGET_ROLE}}", role).replace("{{RESUME_TEXT}}", text)
    raw = _strip_fences(run_llm(prompt))
    try:
        return json.loads(raw)
    except json.JSONDecodeError as exc:
        raise SystemExit(f"Engine did not return valid JSON: {exc}\n--- raw output ---\n{raw}")


def main() -> None:
    ap = argparse.ArgumentParser(description="Parse a résumé file into structured JSON via claude -p.")
    ap.add_argument("file", type=Path, help="résumé file (.pdf .docx .txt .md or an image)")
    ap.add_argument("--role", default="", help="optional target role to weight scoring")
    ap.add_argument("--out", type=Path, help="write JSON here instead of stdout")
    args = ap.parse_args()

    if not args.file.exists():
        raise SystemExit(f"File not found: {args.file}")

    result = parse_resume(args.file, args.role)
    payload = json.dumps(result, indent=2, ensure_ascii=False)

    if args.out:
        args.out.write_text(payload, encoding="utf-8")
        print(f"Wrote {args.out}")
    else:
        print(payload)


if __name__ == "__main__":
    sys.exit(main())
