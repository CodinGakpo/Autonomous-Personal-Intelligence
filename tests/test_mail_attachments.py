"""Tests for the mail pipeline's local-first attachment handling (Excel/PDF)."""

from pathlib import Path

from brain import emailtool as ma_tool
from brain import mail_attachments as ma


def _make_excel(tmp_path, headers, rows):
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.append(headers)
    for row in rows:
        ws.append(row)
    path = tmp_path / "sheet.xlsx"
    wb.save(str(path))
    return path


def test_extract_excel_match_finds_row(tmp_path):
    path = _make_excel(
        tmp_path, ["Name", "Student ID", "Status"],
        [["Asha", "S123", "Selected"], ["Ravi", "S456", "Not Selected"]],
    )
    result = ma.extract_excel_match(path, "S456")
    assert result == {
        "headers": ["Name", "Student ID", "Status"],
        "row": ["Ravi", "S456", "Not Selected"],
    }


def test_extract_excel_match_no_match_returns_none(tmp_path):
    path = _make_excel(tmp_path, ["Name", "Student ID"], [["Asha", "S123"]])
    assert ma.extract_excel_match(path, "S999") is None


def test_extract_excel_match_no_id_column_returns_none(tmp_path):
    path = _make_excel(tmp_path, ["Name", "Score"], [["Asha", "90"]])
    assert ma.extract_excel_match(path, "S123") is None


def test_looks_like_job_description_true_for_jd_text():
    text = "Responsibilities: build stuff. Requirements: Python. Qualifications: CS degree."
    assert ma.looks_like_job_description(text)


def test_looks_like_job_description_false_for_plain_text():
    assert not ma.looks_like_job_description("Hey, lunch at noon?")


def test_extract_pdf_text_caps_length(monkeypatch):
    class FakePage:
        def extract_text(self):
            return "x" * 10000

    class FakeReader:
        def __init__(self, _path):
            self.pages = [FakePage()]

    monkeypatch.setattr("pypdf.PdfReader", FakeReader)
    text = ma.extract_pdf_text(Path("fake.pdf"), max_chars=100)
    assert len(text) == 100


def test_match_resume_to_jd_calls_openrouter(monkeypatch):
    captured = {}

    def fake_call(prompt):
        captured["prompt"] = prompt
        return "Good match on Python."

    monkeypatch.setattr(ma, "call_openrouter", fake_call)
    profile = {"skills": ["Python"], "summary": "backend eng"}
    note = ma.match_resume_to_jd("Need Python dev", profile)
    assert note == "Good match on Python."
    assert "Python" in captured["prompt"]


def test_process_attachment_excel_dispatch(tmp_path):
    path = _make_excel(tmp_path, ["Name", "Student ID"], [["Asha", "S1"]])
    result = ma.process_attachment(path, {"student_id": "S1"})
    assert result["kind"] == "excel"
    # The finding is now a list of matched rows (a sheet can name you more than once) and
    # records which identifier hit, so Q&A can say *why* it thinks the row is yours.
    assert result["finding"] == [
        {"matched": ["S1"], "values": {"Name": "Asha", "Student ID": "S1"}}
    ]
    assert result["mentions_you"] == ["S1"]


def test_process_attachment_unsupported_type(tmp_path):
    # .docx used to land here; it is read now, so this needs a genuinely unhandled type.
    path = tmp_path / "archive.zip"
    path.write_text("hi")
    result = ma.process_attachment(path, {})
    assert result["finding"] == "attachment not parsed"


def test_process_attachment_pdf_non_jd(tmp_path, monkeypatch):
    monkeypatch.setattr(
        ma, "extract_pdf_text",
        lambda path, max_chars=ma.PDF_CHAR_CAP: "Hey, see you soon.",
    )
    path = tmp_path / "note.pdf"
    path.write_bytes(b"%PDF-1.4")
    result = ma.process_attachment(path, {})
    assert result["kind"] == "pdf"


def test_process_attachment_pdf_jd_with_resume(tmp_path, monkeypatch):
    jd_text = "Responsibilities: code. Requirements: Python. Qualifications: BS."
    monkeypatch.setattr(
        ma, "extract_pdf_text",
        lambda path, max_chars=ma.PDF_CHAR_CAP: jd_text,
    )
    monkeypatch.setattr(ma, "match_resume_to_jd", lambda jd, profile: "matches well")
    path = tmp_path / "jd.pdf"
    path.write_bytes(b"%PDF-1.4")
    result = ma.process_attachment(path, {"resume_profile": {"skills": ["Python"]}})
    assert result == {
        "file": "jd.pdf", "kind": "pdf_jd", "finding": "matches well",
        # A JD's raw text is replaced by the match note, so identifier hits ride separately.
        "mentions_you": [],
    }


# --- Identifier matching in attachments -------------------------------------------------


def _sheet(tmp_path, rows, name="shortlist.xlsx"):
    """A real .xlsx on disk — openpyxl parsing is the thing under test."""
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    for row in rows:
        ws.append(row)
    path = tmp_path / name
    wb.save(str(path))
    return path


def test_identifiers_from_config_folds_in_student_id():
    cfg = {"student_id": "23BCE1234", "identifiers": ["Adidev Anand"]}
    assert ma.identifiers_from_config(cfg) == ["Adidev Anand", "23BCE1234"]


def test_identifiers_from_config_ignores_blanks():
    assert ma.identifiers_from_config({"student_id": "  ", "identifiers": []}) == []


def test_find_identifiers_matches_whole_words_only():
    text = "Shortlisted: 23BCE1234, Adidev Anand"
    assert ma.find_identifiers(text, ["23BCE1234"]) == ["23BCE1234"]
    assert ma.find_identifiers(text, ["Adidev Anand"]) == ["Adidev Anand"]
    # A longer id that merely contains yours must not count as you.
    assert ma.find_identifiers("23BCE12345 only", ["23BCE1234"]) == []


def test_find_identifiers_is_case_and_punctuation_insensitive():
    assert ma.find_identifiers("adidev  anand.", ["Adidev Anand"]) == ["Adidev Anand"]


def test_scan_excel_finds_the_row_naming_you(tmp_path):
    path = _sheet(
        tmp_path,
        [
            ["Reg No", "Name", "Company"],
            ["23BCE0001", "Someone Else", "Acme"],
            ["23BCE1234", "Adidev Anand", "Hevo"],
        ],
    )
    scan = ma.scan_excel(path, ["23BCE1234"])
    assert len(scan["matched"]) == 1
    assert scan["matched"][0]["values"]["Name"] == "Adidev Anand"
    assert scan["row_count"] == 2


def test_scan_excel_matches_a_name_in_any_column(tmp_path):
    """The old code only looked at a column whose header contained "id"."""
    path = _sheet(tmp_path, [["Sr", "Candidate"], [1, "Adidev Anand"]])
    scan = ma.scan_excel(path, ["Adidev Anand"])
    assert scan["matched"][0]["matched"] == ["Adidev Anand"]


def test_scan_excel_reports_nothing_when_absent(tmp_path):
    path = _sheet(tmp_path, [["Reg No", "Name"], ["23BCE0001", "Someone Else"]])
    assert ma.scan_excel(path, ["23BCE1234"])["matched"] == []


def test_process_attachment_flags_a_spreadsheet_that_names_you(tmp_path):
    path = _sheet(tmp_path, [["Reg No", "Name"], ["23BCE1234", "Adidev Anand"]])
    finding = ma.process_attachment(path, {"identifiers": ["23BCE1234"]})
    assert finding["kind"] == "excel"
    assert finding["mentions_you"] == ["23BCE1234"]


def test_process_attachment_says_so_when_you_are_not_listed(tmp_path):
    path = _sheet(tmp_path, [["Reg No"], ["23BCE0001"]])
    finding = ma.process_attachment(path, {"identifiers": ["23BCE1234"]})
    assert finding["mentions_you"] == []
    assert "not listed" in finding["finding"]


def test_process_attachment_without_identifiers_still_describes_the_sheet(tmp_path):
    path = _sheet(tmp_path, [["Reg No"], ["23BCE0001"]])
    finding = ma.process_attachment(path, {})
    assert "no identifiers configured" in finding["finding"]


# --- Mailbox status is a real check, not a file-existence placeholder --------------------


def test_mailbox_status_reports_disconnected_without_a_token(tmp_path, monkeypatch):
    monkeypatch.setenv("GMAIL_TOKEN_PATH", str(tmp_path / "token.json"))
    status = ma_tool.mailbox_status(1)
    assert status["connected"] is False
    assert "no mailbox connected" in status["reason"]


def test_mailbox_status_rejects_a_token_file_that_is_not_a_real_grant(tmp_path, monkeypatch):
    """A file can exist and be worthless — that is exactly how a UI ends up claiming a
    connection that fails on the next fetch."""
    token = tmp_path / "token.json"
    token.write_text('{"token": "not-a-real-grant"}', encoding="utf-8")
    monkeypatch.setenv("GMAIL_TOKEN_PATH", str(token))

    status = ma_tool.mailbox_status(1)
    assert status["connected"] is False
    assert status["reason"]


def test_mailbox_status_reports_connected_for_usable_credentials(tmp_path, monkeypatch):
    token = tmp_path / "token.json"
    token.write_text("{}", encoding="utf-8")
    monkeypatch.setenv("GMAIL_TOKEN_PATH", str(token))

    class _Creds:
        valid = True

    monkeypatch.setattr(ma_tool, "_load_oauth_credentials", lambda interactive, user_id: _Creds())
    monkeypatch.setattr(ma_tool, "_mailbox_email", lambda user_id: "me@example.com")

    assert ma_tool.mailbox_status(1) == {"connected": True, "email": "me@example.com"}


def test_mailbox_status_reports_disconnected_when_refresh_fails(tmp_path, monkeypatch):
    token = tmp_path / "token.json"
    token.write_text("{}", encoding="utf-8")
    monkeypatch.setenv("GMAIL_TOKEN_PATH", str(token))

    def _boom(interactive, user_id):
        raise RuntimeError("invalid_scope")

    monkeypatch.setattr(ma_tool, "_load_oauth_credentials", _boom)
    status = ma_tool.mailbox_status(1)
    assert status["connected"] is False
    assert "invalid_scope" in status["reason"]


# --- CSV / DOCX / image attachments must not be dead ends --------------------------------


def test_csv_is_read_and_scanned(tmp_path):
    path = tmp_path / "list.csv"
    path.write_text("Neo Id,Name\nA1B2C3D4,Someone\nT7M8V0L5,Adidev Anand\n", encoding="utf-8")
    result = ma.process_attachment(path, {"identifiers": ["T7M8V0L5"]})
    assert result["kind"] == "csv"
    assert result["mentions_you"] == ["T7M8V0L5"]
    assert "Adidev Anand" in result["finding"]


def test_tsv_delimiter_is_handled(tmp_path):
    path = tmp_path / "list.tsv"
    path.write_text("Neo Id\tName\nT7M8V0L5\tAdidev\n", encoding="utf-8")
    assert ma.process_attachment(path, {"identifiers": ["T7M8V0L5"]})["mentions_you"] == [
        "T7M8V0L5"
    ]


def test_docx_paragraphs_and_tables_are_read(tmp_path):
    from docx import Document

    document = Document()
    document.add_paragraph("Shortlisted candidates for the drive")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "T7M8V0L5"
    table.rows[0].cells[1].text = "Adidev Anand"
    path = tmp_path / "shortlist.docx"
    document.save(str(path))

    result = ma.process_attachment(path, {"identifiers": ["T7M8V0L5"]})
    assert result["kind"] == "docx"
    assert result["mentions_you"] == ["T7M8V0L5"]
    assert "Shortlisted candidates" in result["finding"]


def test_image_text_is_extracted_by_local_ocr(tmp_path, monkeypatch):
    """OCR runs locally first — the image is personal mail and should not leave the machine
    when it does not have to."""
    path = tmp_path / "notice.png"
    path.write_bytes(b"\x89PNG\r\n")
    monkeypatch.setattr(
        ma, "_ocr_with_tesseract", lambda p: "BADA KHANA MENU\nMushroom Biryani\nGhee Phulka"
    )

    def _no_network(p):
        raise AssertionError("vision model must not be called when OCR succeeded")

    monkeypatch.setattr(ma, "_read_image_with_vision_model", _no_network)

    result = ma.process_attachment(path, {"identifiers": []})
    assert result["kind"] == "image"
    assert "Mushroom Biryani" in result["finding"]


def test_image_falls_back_to_a_vision_model_without_tesseract(tmp_path, monkeypatch):
    path = tmp_path / "notice.png"
    path.write_bytes(b"\x89PNG\r\n")
    monkeypatch.setattr(ma, "_ocr_with_tesseract", lambda p: "")
    monkeypatch.setattr(ma, "_read_image_with_vision_model", lambda p: "Menu: Kashmeri Pulav")

    result = ma.process_attachment(path, {"identifiers": []})
    assert result["finding"] == "Menu: Kashmeri Pulav"


def test_speckle_from_a_photo_is_not_mistaken_for_text(tmp_path, monkeypatch):
    path = tmp_path / "photo.png"
    path.write_bytes(b"\x89PNG\r\n")
    monkeypatch.setattr(ma, "_ocr_with_tesseract", lambda p: ". , ' -")
    monkeypatch.setattr(ma, "_read_image_with_vision_model", lambda p: "A group photo.")
    assert ma.process_attachment(path, {"identifiers": []})["finding"] == "A group photo."


def test_an_image_with_nothing_readable_says_so(tmp_path, monkeypatch):
    path = tmp_path / "blank.png"
    path.write_bytes(b"\x89PNG\r\n")
    monkeypatch.setattr(ma, "_ocr_with_tesseract", lambda p: "")
    monkeypatch.setattr(ma, "_read_image_with_vision_model", lambda p: "")
    result = ma.process_attachment(path, {"identifiers": []})
    assert "no readable text" in result["finding"]


def test_a_broken_attachment_does_not_abort_the_run(tmp_path):
    path = tmp_path / "corrupt.docx"
    path.write_bytes(b"not really a docx")
    result = ma.process_attachment(path, {"identifiers": []})
    assert result["kind"] == "docx"
    assert "could not be read" in result["finding"]
    assert result["mentions_you"] == []


def test_an_identifier_inside_an_image_is_found(tmp_path, monkeypatch):
    path = tmp_path / "shortlist.png"
    path.write_bytes(b"\x89PNG\r\n")
    monkeypatch.setattr(ma, "_ocr_with_tesseract", lambda p: "Shortlist: T7M8V0L5, A1B2C3D4")
    assert ma.process_attachment(path, {"identifiers": ["T7M8V0L5"]})["mentions_you"] == [
        "T7M8V0L5"
    ]
