"""brain/mail_attachments.py — local-first attachment handling for the mail pipeline.

Excel: parsed entirely locally (no LLM) — scan every cell for the user's identifiers (name,
roll/neo id) and return the rows that mention them.
PDF: text extracted locally, hard-capped, then (only if it looks like a JD) compared to the
student's résumé profile via one small LLM call.

The local-first, hard-capped design is the guard against a single large attachment burning
through a free-tier OpenRouter key on one email.
"""

from __future__ import annotations

import base64
import csv
import io
import mimetypes
import shutil
import subprocess
from pathlib import Path
from typing import Any

from brain import bm25
from brain.openrouter import call_openrouter, call_openrouter_vision

PDF_CHAR_CAP = 4000
TEXT_CHAR_CAP = 4000
EXCEL_SUFFIXES = {".xlsx", ".xls"}
CSV_SUFFIXES = {".csv", ".tsv"}
DOCX_SUFFIXES = {".docx"}
IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tif", ".tiff"}
MAX_MATCHED_ROWS = 5
OCR_TIMEOUT_S = 90
# Notices, menus and posters routinely arrive as a photo or screenshot with the actual content
# baked into pixels. Leaving those unread makes the mail unanswerable, so images get read too.
VISION_MODEL_ENV = "OPENROUTER_VISION_MODEL"
DEFAULT_VISION_MODEL = "meta-llama/llama-3.2-11b-vision-instruct:free"
JD_KEYWORDS = (
    "responsibilities", "requirements", "qualifications",
    "job description", "role overview", "skills required",
)


def identifiers_from_config(config: dict[str, Any]) -> list[str]:
    """Everything that counts as "me" in an attachment: name, roll/neo id, registration no.

    `student_id` stays supported so existing configs keep working; `identifiers` is the general
    form, because a shortlist is just as likely to list a name as an id.
    """
    values = list(config.get("identifiers") or [])
    single = (config.get("student_id") or "").strip()
    if single and single not in values:
        values.append(single)
    return [v for v in (str(v).strip() for v in values) if v]


def _zone(text: str) -> str:
    """Normalized, space-padded text so identifier matching is whole-word, not substring."""
    return " " + " ".join(bm25.tokenize(text or "")) + " "


def find_identifiers(text: str, identifiers: list[str]) -> list[str]:
    """Which identifiers appear in `text`, matched as whole words (case/punctuation-insensitive).

    Whole-word matching matters both ways: "23BCE1234" must not match "23BCE12345", and a name
    must not match a fragment of another word.
    """
    zone = _zone(text)
    hits = []
    for identifier in identifiers:
        tokens = bm25.tokenize(identifier)
        if tokens and f" {' '.join(tokens)} " in zone:
            hits.append(identifier)
    return hits


def extract_excel_match(path: Path, student_id: str) -> dict | None:
    """Find the row whose ID-like column matches `student_id`. None if no match or no ID column.

    Kept for the narrow exact-id case; `scan_excel` is the general entry point.
    """
    from openpyxl import load_workbook

    wb = load_workbook(str(path), read_only=True, data_only=True)
    sheet = wb.active
    rows = list(sheet.iter_rows(values_only=True))
    if not rows:
        return None
    headers = [str(h) if h is not None else "" for h in rows[0]]
    id_col = next((i for i, h in enumerate(headers) if "id" in h.lower()), None)
    if id_col is None:
        return None
    for row in rows[1:]:
        cell = row[id_col]
        if cell is not None and str(cell).strip() == str(student_id).strip():
            return {"headers": headers, "row": [str(c) if c is not None else "" for c in row]}
    return None


def scan_excel(path: Path, identifiers: list[str]) -> dict[str, Any]:
    """Search every cell of the sheet for any identifier, returning the rows that mention one.

    Deliberately not restricted to a column whose header contains "id": real shortlists put the
    name in one column and the roll number in another, spell the header a dozen ways, or have
    no header row at all. Scanning cells is what makes "is my name in this list?" answerable.
    """
    from openpyxl import load_workbook

    wb = load_workbook(str(path), read_only=True, data_only=True)
    sheet = wb.active
    rows = list(sheet.iter_rows(values_only=True))
    if not rows:
        return {"matched": [], "row_count": 0, "headers": []}

    headers = [str(h) if h is not None else "" for h in rows[0]]
    matched: list[dict[str, Any]] = []
    for row in rows[1:]:
        cells = [str(c) if c is not None else "" for c in row]
        hits = find_identifiers(" ".join(cells), identifiers)
        if not hits:
            continue
        padded = list(headers) + [""] * max(0, len(cells) - len(headers))
        labelled = {
            h or f"col{i + 1}": c
            for i, (h, c) in enumerate(zip(padded, cells, strict=False))
        }
        matched.append({"matched": hits, "values": labelled})
        if len(matched) >= MAX_MATCHED_ROWS:
            break
    return {"matched": matched, "row_count": max(len(rows) - 1, 0), "headers": headers}


def extract_pdf_text(path: Path, max_chars: int = PDF_CHAR_CAP) -> str:
    """Extract PDF text via pypdf, hard-capped at `max_chars`."""
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    text = "\n".join(page.extract_text() or "" for page in reader.pages)
    return text[:max_chars]


def extract_csv_text(path: Path, max_chars: int = TEXT_CHAR_CAP) -> str:
    """Read a delimited file as text, sniffing the delimiter (comma vs tab)."""
    raw = path.read_text(encoding="utf-8", errors="replace")
    delimiter = "\t" if path.suffix.lower() == ".tsv" or "\t" in raw[:400] else ","
    rows = list(csv.reader(io.StringIO(raw), delimiter=delimiter))
    return "\n".join(", ".join(cell for cell in row if cell) for row in rows)[:max_chars]


def extract_docx_text(path: Path, max_chars: int = TEXT_CHAR_CAP) -> str:
    """Paragraph and table text from a .docx (python-docx is already a project dependency)."""
    from docx import Document

    document = Document(str(path))
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)[:max_chars]


def extract_image_text(path: Path, max_chars: int = TEXT_CHAR_CAP) -> str:
    """Text inside an image, by local OCR then a vision model.

    Local first: Tesseract is free, fast, and keeps the image on this machine — which matters,
    because these are someone's personal mail attachments. The hosted vision model is only a
    fallback for machines with no Tesseract installed, so the feature still works there.
    """
    text = _ocr_with_tesseract(path)
    if _is_meaningful(text):
        return text[:max_chars]
    return _read_image_with_vision_model(path)[:max_chars]


def _ocr_with_tesseract(path: Path) -> str:
    """Shell out to the Tesseract binary — no Python wrapper dependency needed."""
    if shutil.which("tesseract") is None:
        return ""
    try:
        proc = subprocess.run(
            ["tesseract", str(path), "-"],
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
            timeout=OCR_TIMEOUT_S,
        )
    except (OSError, subprocess.SubprocessError):
        return ""
    return proc.stdout.strip() if proc.returncode == 0 else ""


def _is_meaningful(text: str) -> bool:
    """Whether OCR produced something worth keeping rather than speckle from a photo."""
    letters = sum(c.isalnum() for c in text)
    return letters >= 24


def _read_image_with_vision_model(path: Path) -> str:
    """Fallback when Tesseract isn't available: ask a vision model what the image says."""
    try:
        data_url = _data_url(path)
    except OSError:
        return ""
    prompt = (
        "Transcribe every piece of text in this image, preserving structure (dates, times, "
        "venues, menus, lists, tables). If it is a notice or poster, keep the details exactly. "
        "Reply with the text only, no commentary."
    )
    try:
        return call_openrouter_vision(prompt, data_url).strip()
    except (Exception, SystemExit):
        # No keys, no vision model, or the call failed — the caller reports it unread rather
        # than taking the whole ingest down over one attachment.
        return ""


def _data_url(path: Path) -> str:
    mime = mimetypes.guess_type(str(path))[0] or "image/png"
    return f"data:{mime};base64,{base64.b64encode(path.read_bytes()).decode('ascii')}"


def looks_like_job_description(text: str) -> bool:
    """Cheap local heuristic — no LLM call needed just to decide whether to bother matching."""
    lowered = text.lower()
    return sum(1 for kw in JD_KEYWORDS if kw in lowered) >= 2


def match_resume_to_jd(jd_text: str, resume_profile: dict) -> str:
    """One small LLM call: does this student's résumé match this JD? Short match/gap note."""
    skills = ", ".join(resume_profile.get("skills", []))
    summary = resume_profile.get('summary') or resume_profile.get('headline') or ''
    prompt = (
        "You are screening ONE candidate against ONE job description. Reply with 2-3 short "
        "sentences: how well they match, and the clearest gap if any. No preamble, no markdown.\n\n"
        f"CANDIDATE SKILLS: {skills}\n"
        f"CANDIDATE SUMMARY: {summary}\n\n"
        f"JOB DESCRIPTION:\n{jd_text}"
    )
    return call_openrouter(prompt).strip()


def process_attachment(path: Path, config: dict) -> dict:
    """Dispatch one saved attachment file by suffix. Always returns a small, prompt-safe finding."""
    suffix = Path(path).suffix.lower()

    identifiers = identifiers_from_config(config)

    if suffix in EXCEL_SUFFIXES:
        scan = scan_excel(path, identifiers)
        if not identifiers:
            finding = f"spreadsheet with {scan['row_count']} rows (no identifiers configured)"
        elif scan["matched"]:
            finding = scan["matched"]
        else:
            finding = f"you were not listed in this spreadsheet ({scan['row_count']} rows)"
        return {
            "file": Path(path).name,
            "kind": "excel",
            "finding": finding,
            "mentions_you": [m for row in scan["matched"] for m in row["matched"]],
        }

    if suffix == ".pdf":
        text = extract_pdf_text(path)
        mentions = find_identifiers(text, identifiers)
        if not looks_like_job_description(text):
            return {
                "file": Path(path).name,
                "kind": "pdf",
                "finding": text[:500],
                "mentions_you": mentions,
            }
        resume_profile = config.get("resume_profile") or {}
        if resume_profile:
            note = match_resume_to_jd(text, resume_profile)
        else:
            note = "no résumé configured to match against"
        return {
            "file": Path(path).name,
            "kind": "pdf_jd",
            "finding": note,
            # A JD's text is replaced by the match note, so carry the hit separately or the
            # fact that you were named in it would be lost.
            "mentions_you": mentions,
        }

    # Notices, menus and posters routinely arrive as an image or a plain table. Leaving these
    # unread makes the mail unanswerable ("what's on the mess menu?"), so they are read too.
    readers: list[tuple[set[str], str, Any]] = [
        (CSV_SUFFIXES, "csv", extract_csv_text),
        (DOCX_SUFFIXES, "docx", extract_docx_text),
        (IMAGE_SUFFIXES, "image", extract_image_text),
    ]
    for suffixes, kind, extract in readers:
        if suffix not in suffixes:
            continue
        try:
            text = extract(path)
        except Exception as exc:  # a malformed attachment must not take the whole ingest down
            return {
                "file": Path(path).name,
                "kind": kind,
                "finding": f"could not be read: {exc}",
                "mentions_you": [],
            }
        if not text.strip():
            return {
                "file": Path(path).name,
                "kind": kind,
                "finding": "no readable text found in this file",
                "mentions_you": [],
            }
        return {
            "file": Path(path).name,
            "kind": kind,
            "finding": text,
            "mentions_you": find_identifiers(text, identifiers),
        }

    return {
        "file": Path(path).name,
        "kind": suffix.lstrip(".") or "unknown",
        "finding": "attachment not parsed",
        "mentions_you": [],
    }
